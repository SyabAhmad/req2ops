import io
import base64
from fastapi import UploadFile, HTTPException
from app.config import settings


AUDIO_EXTENSIONS = (".mp3", ".wav", ".m4a", ".aac", ".ogg", ".wma", ".flac", ".webm")
VIDEO_EXTENSIONS = (".mp4", ".mkv", ".avi", ".mov", ".flv", ".wmv")
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg")
DOCX_EXTENSIONS = (".docx", ".doc")
CSV_EXTENSIONS = (".csv", ".tsv")
XLSX_EXTENSIONS = (".xlsx", ".xls")
TEXT_EXTENSIONS = (".txt", ".md", ".json", ".xml")


class ParserService:
    def __init__(self):
        self.allowed_extensions = tuple(
            ext.strip() for ext in settings.allowed_file_extensions.split(",") if ext.strip()
        )
        self.max_bytes = settings.max_file_size_mb * 1024 * 1024

    async def parse_file(self, file: UploadFile) -> str:
        content = await file.read()
        if len(content) > self.max_bytes:
            raise HTTPException(
                status_code=413,
                detail=f"File exceeds max size of {settings.max_file_size_mb}MB",
            )

        filename = file.filename or "unknown"
        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext not in self.allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{ext}'. Allowed: {settings.allowed_file_extensions}",
            )

        if ext == ".pdf":
            return self._parse_pdf(content, filename)

        if ext in AUDIO_EXTENSIONS:
            return await self._transcribe_audio(content, filename)

        if ext in VIDEO_EXTENSIONS:
            return await self._transcribe_video(content, filename)

        if ext in IMAGE_EXTENSIONS:
            return await self._ocr_image(content, filename)

        if ext in DOCX_EXTENSIONS:
            return self._parse_docx(content, filename)

        if ext in CSV_EXTENSIONS:
            return self._parse_csv(content, filename)

        if ext in XLSX_EXTENSIONS:
            return self._parse_xlsx(content, filename)

        if ext in TEXT_EXTENSIONS:
            return content.decode("utf-8", errors="replace")

        return content.decode("utf-8", errors="replace")

    def parse_text(self, text: str) -> str:
        return text.strip()

    def _parse_pdf(self, content: bytes, filename: str) -> str:
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return f"[PDF Extracted — {filename}]\n{text}"
        except ImportError:
            return content.decode("utf-8", errors="replace")

    def _parse_docx(self, content: bytes, filename: str) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables_text.append(" | ".join(cells))
            if tables_text:
                text += "\n\n[Tables]\n" + "\n".join(tables_text)

            return f"[DOCX Extracted — {filename}]\n{text}"
        except ImportError:
            return f"[DOCX file uploaded: {filename}] — python-docx not installed"

    def _parse_csv(self, content: bytes, filename: str) -> str:
        try:
            import csv
            text = content.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
            if not rows:
                return f"[CSV file empty: {filename}]"

            headers = rows[0]
            data_rows = rows[1:51]

            lines = [f"[CSV Extracted — {filename}]"]
            lines.append("Headers: " + " | ".join(headers))
            lines.append(f"Total rows: {len(rows) - 1}")
            lines.append("")
            for i, row in enumerate(data_rows, 1):
                lines.append(f"Row {i}: " + " | ".join(row))
            if len(rows) > 51:
                lines.append(f"\n... and {len(rows) - 51} more rows")

            return "\n".join(lines)
        except Exception:
            return f"[CSV file uploaded: {filename}]"

    def _parse_xlsx(self, content: bytes, filename: str) -> str:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines = [f"[XLSX Extracted — {filename}]"]

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"\nSheet: {sheet_name}")
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    lines.append("  (empty)")
                    continue

                headers = [str(c) if c else "" for c in rows[0]]
                lines.append("Headers: " + " | ".join(headers))

                for i, row in enumerate(rows[1:51], 1):
                    cells = [str(c) if c else "" for c in row]
                    lines.append(f"Row {i}: " + " | ".join(cells))
                if len(rows) > 51:
                    lines.append(f"  ... and {len(rows) - 51} more rows")

            wb.close()
            return "\n".join(lines)
        except ImportError:
            return f"[XLSX file uploaded: {filename}] — openpyxl not installed"
        except Exception:
            return f"[XLSX file uploaded: {filename}]"

    async def _ocr_image(self, content: bytes, filename: str) -> str:
        try:
            from groq import Groq
            client = Groq(api_key=settings.groq_api_key)

            b64 = base64.b64encode(content).decode("utf-8")
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
            mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp"}
            mime = mime_map.get(ext, "image/png")

            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{mime};base64,{b64}"},
                            },
                            {
                                "type": "text",
                                "text": "Extract all text from this image. If it's a screenshot of a document, chat, or notes, transcribe all visible text. Return only the extracted text, no commentary.",
                            },
                        ],
                    }
                ],
                max_tokens=4096,
            )
            text = response.choices[0].message.content
            return f"[Image OCR — {filename}]\n{text}"
        except Exception:
            return f"[Image uploaded: {filename}] — OCR failed"

    async def _transcribe_audio(self, content: bytes, filename: str) -> str:
        from app.services.transcription import TranscriptionService
        ts = TranscriptionService()
        text = await ts.transcribe_audio(content, filename)
        return f"[Audio Transcription — {filename}]\n{text}"

    async def _transcribe_video(self, content: bytes, filename: str) -> str:
        from app.services.transcription import TranscriptionService
        ts = TranscriptionService()
        text = await ts.transcribe_video(content, filename)
        return f"[Video Transcription — {filename}]\n{text}"
